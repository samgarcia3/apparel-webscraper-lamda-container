from docx import Document  #For Word document
import csv  # For CSV
from bs4 import BeautifulSoup
import requests 
 
 
url_to_parse = "https://en.wikipedia.org/wiki/Python_(programming_language)"
 
response = requests.get(url_to_parse)
response_text = response.text
 
soup = BeautifulSoup(response_text, 'lxml')
heading_strings = []
main_headings = soup.find_all("li", class_="toclevel-1")
for h1 in main_headings:
    heading = h1.find("span", class_="toctext").text
    heading_strings.append(heading)
 
# Write to word Document
document = Document()
document.add_paragraph(f'This is table of contents from {url_to_parse}')
 
for heading in heading_strings:
    document.add_paragraph(heading)
 
document.save('TOC Level 1.docx')
# Write to CSV
with open("toc level 1.csv", "w", newline="") as file:
    writer = csv.writer(file)
    for heading in heading_strings:
        writer.writerow([heading])  # Convert string to list